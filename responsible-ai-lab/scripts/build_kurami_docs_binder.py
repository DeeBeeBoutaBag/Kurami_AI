from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "generated_docs"
OUT_PATH = OUT_DIR / "Kurami_AI_Production_Binder.docx"

SOURCE_FILES = [
    "README.md",
    "docs/codes-and-steps.md",
    "docs/facilitator-guide.md",
    "docs/workshop-explainers.md",
    "docs/manual-fallback-guide.md",
    "docs/event-day-checklist.md",
    "docs/production-risks.md",
    "docs/privacy-and-safety.md",
    "docs/architecture-summary.md",
    "docs/architecture.md",
    "docs/socket-events.md",
    "docs/load-testing.md",
    "docs/implementation-checklist.md",
]


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def set_cell_width(cell, width_in: float):
    width = int(width_in * 1440)
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_layout_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 20, "000000", 20, 6),
        ("Heading 2", 16, "000000", 18, 6),
        ("Heading 3", 14, "434343", 16, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for list_style_name in ["List Bullet", "List Number"]:
        style = styles[list_style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    if "Source Path" not in styles:
        source = styles.add_style("Source Path", 1)
    else:
        source = styles["Source Path"]
    source.font.name = "Arial"
    source._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    source.font.size = Pt(9)
    source.font.color.rgb = RGBColor(85, 85, 85)
    source.paragraph_format.space_before = Pt(0)
    source.paragraph_format.space_after = Pt(10)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(0, 0, 0)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Kurami.AI Production Binder")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Compiled from README.md and all repository Markdown documentation.")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("Contents", level=1)
    for src in SOURCE_FILES:
        doc.add_paragraph(src, style="List Bullet")


def strip_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("*", "")
    return text.strip()


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")


def add_inline_runs(paragraph, text: str):
    for part in filter(None, INLINE_RE.split(text)):
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run.font.size = Pt(10)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("["):
            match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if match:
                paragraph.add_run(f"{match.group(1)} ({match.group(2)})")
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)


def parse_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [strip_inline(cell) for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    header = rows[0]
    body = rows[2:] if len(rows) > 1 and all(set(c.replace(" ", "")) <= {"-", ":"} for c in rows[1]) else rows[1:]
    width = max(len(header), *(len(row) for row in body)) if body else len(header)
    header += [""] * (width - len(header))
    fixed_body = []
    for row in body:
        fixed_body.append(row + [""] * (width - len(row)))
    return header, fixed_body


def add_table(doc: Document, lines: list[str]):
    header, body = parse_table(lines)
    cols = len(header)
    if cols == 0:
        return
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout_fixed(table)
    set_table_borders(table)

    total_width = 6.5
    if cols == 2:
        widths = [2.0, 4.5]
    elif cols == 3:
        widths = [1.5, 2.0, 3.0]
    elif cols == 4:
        widths = [1.35, 1.7, 1.55, 1.9]
    elif cols == 5:
        widths = [1.1, 1.25, 1.25, 1.35, 1.55]
    elif cols == 6:
        widths = [1.0, 1.1, 1.0, 1.25, 1.15, 1.0]
    else:
        widths = [total_width / cols] * cols

    for index, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[index])
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header[index])
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)

    for row in body:
        cells = table.add_row().cells
        for index, cell in enumerate(cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, row[index])
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)

    doc.add_paragraph()


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return lines[index].lstrip().startswith("|") and lines[index + 1].lstrip().startswith("|") and "---" in lines[index + 1]


def collect_table(lines: list[str], index: int) -> tuple[list[str], int]:
    block = []
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        block.append(lines[index])
        index += 1
    return block, index


def collect_paragraph(lines: list[str], index: int) -> tuple[str, int]:
    parts = []
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            break
        if stripped.startswith("#") or stripped.startswith("```") or is_table_start(lines, index):
            break
        if re.match(r"^\s*([-*]|\d+\.)\s+", line):
            break
        parts.append(stripped)
        index += 1
    return " ".join(parts), index


def add_markdown_file(doc: Document, source: str, first: bool):
    path = ROOT / source
    if not first:
        doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_paragraph(source, style="Source Path")
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_language = ""
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = not in_code
            code_language = stripped[3:].strip()
            if in_code and code_language:
                p = doc.add_paragraph(style="Code Block")
                add_inline_runs(p, f"{code_language}")
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(style="Code Block")
            p.add_run(line if line else " ")
            i += 1
            continue

        if is_table_start(lines, i):
            table_lines, i = collect_table(lines, i)
            add_table(doc, table_lines)
            continue

        header = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if header:
            level = min(len(header.group(1)), 3)
            p = doc.add_heading("", level=level)
            add_inline_runs(p, header.group(2))
            i += 1
            continue

        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if list_match:
            marker = list_match.group(2)
            text = list_match.group(3)
            indent_spaces = len(list_match.group(1))
            if marker.endswith("."):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.42 + min(indent_spaces // 2, 3) * 0.18)
                p.paragraph_format.first_line_indent = Inches(-0.22)
                p.paragraph_format.space_after = Pt(4)
                p.add_run(f"{marker} ")
                add_inline_runs(p, text)
            else:
                p = doc.add_paragraph(style="List Bullet")
                if indent_spaces >= 2:
                    p.paragraph_format.left_indent = Inches(0.25 + min(indent_spaces // 2, 3) * 0.18)
                add_inline_runs(p, text)
            i += 1
            continue

        paragraph_text, new_index = collect_paragraph(lines, i)
        if paragraph_text:
            p = doc.add_paragraph()
            add_inline_runs(p, paragraph_text)
            i = new_index
        else:
            i += 1


def build():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title(doc)
    first = False
    for src in SOURCE_FILES:
        add_markdown_file(doc, src, first=first)
        first = False
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
